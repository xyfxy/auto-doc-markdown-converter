#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
创建测试 DOCX 文档的脚本
"""

from docx import Document

def create_test_docx():
    """创建一个简单的测试 DOCX 文档"""
    doc = Document()
    
    # 添加主标题
    doc.add_heading('智能文档处理测试', 0)
    
    # 添加介绍段落
    doc.add_paragraph('这是一个用于测试智能文档处理器的示例文档。')
    
    # 添加一级标题
    doc.add_heading('第一章：基本功能', level=1)
    doc.add_paragraph('这个章节将介绍基本功能。')
    
    # 添加二级标题
    doc.add_heading('1.1 文档解析', level=2)
    doc.add_paragraph('文档解析是核心功能之一，可以智能识别文档结构。')
    
    # 添加三级标题
    doc.add_heading('1.1.1 文本提取', level=3)
    doc.add_paragraph('从各种格式的文档中提取纯文本内容。')
    
    # 添加另一个一级标题
    doc.add_heading('第二章：高级功能', level=1)
    doc.add_paragraph('这个章节将介绍高级功能和配置选项。')
    
    # 保存文档
    doc.save('real_test.docx')
    print("测试文档 'real_test.docx' 创建成功！")

if __name__ == "__main__":
    create_test_docx() 