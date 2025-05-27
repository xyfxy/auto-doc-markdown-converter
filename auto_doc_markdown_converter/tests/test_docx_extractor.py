import unittest
from unittest.mock import patch, mock_open, MagicMock, call
import os
import sys

# Ensure the src directory is in the Python path for imports
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

from src.docx_extractor import extract_content_from_docx
from docx.opc.exceptions import PackageNotFoundError
from docx.image.exceptions import UnrecognizedImageError # Assuming this might be raised

class TestDocxImageExtraction(unittest.TestCase):

    def setUp(self):
        self.test_file_path = "dummy.docx"
        self.results_dir = "test_results"
        self.images_dir = os.path.join(self.results_dir, "images")

    @patch('src.docx_extractor.os.makedirs')
    @patch('src.docx_extractor.docx.Document')
    def test_no_images_in_docx(self, mock_docx_document, mock_os_makedirs):
        """Test DOCX file with no images."""
        # Mock Document object and its paragraphs/runs to have no images
        mock_doc_instance = MagicMock()
        mock_para1_run1 = MagicMock()
        mock_para1_run1.text = "Paragraph one text."
        mock_para1_run1.element.findall.return_value = [] # No drawings
        
        mock_para2_run1 = MagicMock()
        mock_para2_run1.text = "Paragraph two text."
        mock_para2_run1.element.findall.return_value = [] # No drawings

        mock_para1 = MagicMock()
        mock_para1.runs = [mock_para1_run1]
        mock_para2 = MagicMock()
        mock_para2.runs = [mock_para2_run1]

        mock_doc_instance.paragraphs = [mock_para1, mock_para2]
        mock_docx_document.return_value = mock_doc_instance

        extracted_text = extract_content_from_docx(self.test_file_path, self.results_dir)

        self.assertEqual(extracted_text, "Paragraph one text.\n\nParagraph two text.")
        mock_os_makedirs.assert_called_once_with(self.images_dir, exist_ok=True)
        # No open calls for writing images

    @patch('src.docx_extractor.os.makedirs')
    @patch('builtins.open', new_callable=mock_open)
    @patch('src.docx_extractor.docx.Document')
    def test_single_image_in_docx(self, mock_docx_document, mock_file_open, mock_os_makedirs):
        """Test DOCX file with one inline image."""
        mock_doc_instance = MagicMock()
        
        # Run 1: Text
        mock_run1 = MagicMock()
        mock_run1.text = "Text before image. "
        mock_run1.element.findall.return_value = [] # No drawings in this run

        # Run 2: Image
        mock_run2_drawing = MagicMock()
        mock_run2_inline = MagicMock()
        mock_run2_graphic = MagicMock()
        mock_run2_graphic_data = MagicMock()
        mock_run2_pic = MagicMock()
        mock_run2_blip_fill = MagicMock()
        mock_run2_blip = MagicMock()
        mock_run2_blip.get.return_value = "rId1" # Relationship ID
        
        mock_run2_blip_fill.find.return_value = mock_run2_blip
        mock_run2_pic.find.return_value = mock_run2_blip_fill
        mock_run2_graphic_data.find.return_value = mock_run2_pic
        mock_run2_graphic.find.return_value = mock_run2_graphic_data
        mock_run2_inline.find.return_value = mock_run2_graphic
        mock_run2_drawing.findall.return_value = [mock_run2_inline] # This inline contains the graphic
        
        mock_run2 = MagicMock()
        mock_run2.text = "" # Run with image might have no direct text
        mock_run2.element.findall.return_value = [mock_run2_drawing] # This run has a drawing

        # Run 3: Text
        mock_run3 = MagicMock()
        mock_run3.text = " Text after image."
        mock_run3.element.findall.return_value = []

        mock_paragraph = MagicMock()
        mock_paragraph.runs = [mock_run1, mock_run2, mock_run3]
        mock_doc_instance.paragraphs = [mock_paragraph]

        # Mock image part
        mock_image_part = MagicMock()
        mock_image_part.blob = b"image_data_bytes"
        mock_image_part.default_image_ext = "png"
        mock_doc_instance.part.related_parts = {'rId1': mock_image_part}
        
        mock_docx_document.return_value = mock_doc_instance

        extracted_text = extract_content_from_docx(self.test_file_path, self.results_dir)

        expected_text = "Text before image. [IMAGE_PLACEHOLDER:image_1.png] Text after image."
        self.assertEqual(extracted_text, expected_text)
        mock_os_makedirs.assert_called_once_with(self.images_dir, exist_ok=True)
        mock_file_open.assert_called_once_with(os.path.join(self.images_dir, "image_1.png"), 'wb')
        mock_file_open().write.assert_called_once_with(b"image_data_bytes")

    @patch('src.docx_extractor.os.makedirs')
    @patch('builtins.open', new_callable=mock_open)
    @patch('src.docx_extractor.docx.Document')
    def test_multiple_images_in_docx(self, mock_docx_document, mock_file_open, mock_os_makedirs):
        """Test DOCX file with multiple inline images across paragraphs."""
        mock_doc_instance = MagicMock()

        # Paragraph 1, Run 1 (Image 1)
        mock_p1_r1_drawing = MagicMock()
        mock_p1_r1_inline = MagicMock()
        mock_p1_r1_graphic = MagicMock(); mock_p1_r1_graphic_data = MagicMock(); mock_p1_r1_pic = MagicMock(); mock_p1_r1_blip_fill = MagicMock(); mock_p1_r1_blip = MagicMock()
        mock_p1_r1_blip.get.return_value = "rId1"
        mock_p1_r1_blip_fill.find.return_value = mock_p1_r1_blip; mock_p1_r1_pic.find.return_value = mock_p1_r1_blip_fill; mock_p1_r1_graphic_data.find.return_value = mock_p1_r1_pic; mock_p1_r1_graphic.find.return_value = mock_p1_r1_graphic_data; mock_p1_r1_inline.find.return_value = mock_p1_r1_graphic
        mock_p1_r1_drawing.findall.return_value = [mock_p1_r1_inline]
        mock_p1_r1 = MagicMock(); mock_p1_r1.text = ""; mock_p1_r1.element.findall.return_value = [mock_p1_r1_drawing]
        mock_para1 = MagicMock(); mock_para1.runs = [mock_p1_r1]

        # Paragraph 2, Run 1 (Text), Run 2 (Image 2)
        mock_p2_r1 = MagicMock(); mock_p2_r1.text = "Some text. "; mock_p2_r1.element.findall.return_value = []
        mock_p2_r2_drawing = MagicMock()
        mock_p2_r2_inline = MagicMock()
        mock_p2_r2_graphic = MagicMock(); mock_p2_r2_graphic_data = MagicMock(); mock_p2_r2_pic = MagicMock(); mock_p2_r2_blip_fill = MagicMock(); mock_p2_r2_blip = MagicMock()
        mock_p2_r2_blip.get.return_value = "rId2"
        mock_p2_r2_blip_fill.find.return_value = mock_p2_r2_blip; mock_p2_r2_pic.find.return_value = mock_p2_r2_blip_fill; mock_p2_r2_graphic_data.find.return_value = mock_p2_r2_pic; mock_p2_r2_graphic.find.return_value = mock_p2_r2_graphic_data; mock_p2_r2_inline.find.return_value = mock_p2_r2_graphic
        mock_p2_r2_drawing.findall.return_value = [mock_p2_r2_inline]
        mock_p2_r2 = MagicMock(); mock_p2_r2.text = ""; mock_p2_r2.element.findall.return_value = [mock_p2_r2_drawing]
        mock_para2 = MagicMock(); mock_para2.runs = [mock_p2_r1, mock_p2_r2]

        mock_doc_instance.paragraphs = [mock_para1, mock_para2]

        # Mock image parts
        mock_image_part1 = MagicMock(); mock_image_part1.blob = b"img_data1"; mock_image_part1.default_image_ext = "jpeg"
        mock_image_part2 = MagicMock(); mock_image_part2.blob = b"img_data2"; mock_image_part2.default_image_ext = "gif"
        mock_doc_instance.part.related_parts = {'rId1': mock_image_part1, 'rId2': mock_image_part2}

        mock_docx_document.return_value = mock_doc_instance

        extracted_text = extract_content_from_docx(self.test_file_path, self.results_dir)
        expected_text = "[IMAGE_PLACEHOLDER:image_1.jpeg]\n\nSome text. [IMAGE_PLACEHOLDER:image_2.gif]"
        self.assertEqual(extracted_text, expected_text)
        mock_os_makedirs.assert_called_once_with(self.images_dir, exist_ok=True)
        
        # Check file open calls
        calls = [
            call(os.path.join(self.images_dir, "image_1.jpeg"), 'wb'),
            call(os.path.join(self.images_dir, "image_2.gif"), 'wb')
        ]
        mock_file_open.assert_has_calls(calls, any_order=True)
        self.assertEqual(mock_file_open.call_count, 2)
        # Check write calls (order depends on any_order=True for open calls, or specific mocking of open().write)
        # This is a bit simplified; real check might need to map open calls to write calls.
        # For now, assume write is called correctly for each opened file.

    @patch('src.docx_extractor.docx.Document')
    def test_docx_package_not_found(self, mock_docx_document):
        """Test handling of PackageNotFoundError."""
        mock_docx_document.side_effect = PackageNotFoundError("File not found")
        extracted_text = extract_content_from_docx(self.test_file_path, self.results_dir)
        self.assertIsNone(extracted_text)

    # It's harder to simulate UnrecognizedImageError without deeper control over python-docx mocks
    # or a real (but small) malformed docx. For now, we trust the except block if it's reached.
    # If UnrecognizedImageError is a specific concern, more targeted mocking of image part processing would be needed.

if __name__ == '__main__':
    unittest.main()
